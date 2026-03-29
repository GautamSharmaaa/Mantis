from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import simpleSplit
from reportlab.platypus import Paragraph
from reportlab.pdfgen import canvas as pdf_canvas

from models.resume_model import Resume
from utils.text_utils import clean_text

_FILENAME_PATTERN = re.compile(r"[^a-zA-Z0-9_-]+")

# ---------------------------------------------------------------------------
# Page geometry constants
# ---------------------------------------------------------------------------
PAGE_W, PAGE_H = LETTER          # 612 × 792 pt


# ===========================================================================
# Shared helpers
# ===========================================================================

def normalize_export_profile(profile: dict[str, str] | None) -> dict[str, str]:
    profile = profile or {}
    return {
        "fullName":  clean_text(profile.get("fullName")),
        "email":     clean_text(profile.get("email")),
        "phone":     clean_text(profile.get("phone")),
        "location":  clean_text(profile.get("location")),
        "jobTitle":  clean_text(profile.get("jobTitle")),
        "website":   clean_text(profile.get("website")),
        "linkedin":  clean_text(profile.get("linkedin")),
        "github":    clean_text(profile.get("github")),
    }


def resume_to_formatted_text(resume: Resume, profile: dict[str, str] | None = None) -> str:
    """Plain-text fallback (ATS-safe)."""
    p = normalize_export_profile(profile)
    display_name = p["fullName"] or resume.title
    sections: list[str] = [display_name]

    contact = " | ".join(v for v in [p["email"], p["phone"], p["location"]] if v)
    if contact:
        sections += [contact, ""]

    if p["jobTitle"]:
        sections += [p["jobTitle"], ""]

    if clean_text(resume.data.summary):
        sections += ["SUMMARY", resume.data.summary, ""]

    if resume.data.experience:
        sections.append("EXPERIENCE")
        for exp in resume.data.experience:
            sections.append(f"{exp.role} | {exp.company}")
            for pt in exp.points:
                sections.append(f"  - {pt}")
            sections.append("")

    if resume.data.projects:
        sections.append("PROJECTS")
        for proj in resume.data.projects:
            sections.append(proj.name)
            for pt in proj.points:
                sections.append(f"  - {pt}")
            sections.append("")

    if resume.data.skills:
        sections += ["SKILLS", ", ".join(resume.data.skills), ""]

    return "\n".join(l for l in sections if l is not None).strip()


def generate_docx_bytes(resume: Resume, profile: dict[str, str] | None = None) -> bytes:
    p = normalize_export_profile(profile)
    doc = Document()

    doc.add_heading(p["fullName"] or resume.title, 0)

    contact = " | ".join(v for v in [p["email"], p["phone"], p["location"]] if v)
    if contact:
        doc.add_paragraph(contact)

    links = " | ".join(
        v.replace("https://", "").replace("http://", "")
        for v in [p["website"], p["github"], p["linkedin"]]
        if v
    )
    if links:
        doc.add_paragraph(links)

    if p["jobTitle"]:
        doc.add_paragraph(p["jobTitle"])

    if clean_text(resume.data.summary):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume.data.summary)

    if resume.data.experience:
        doc.add_heading("Experience", level=1)
        for exp in resume.data.experience:
            doc.add_paragraph(f"{exp.role} – {exp.company}", style="Heading 2")
            for pt in exp.points:
                doc.add_paragraph(pt, style="List Bullet")

    if resume.data.projects:
        doc.add_heading("Projects", level=1)
        for proj in resume.data.projects:
            doc.add_paragraph(proj.name, style="Heading 2")
            for pt in proj.points:
                doc.add_paragraph(pt, style="List Bullet")

    if resume.data.skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(resume.data.skills))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_pdf_bytes(resume: Resume, profile: dict[str, str] | None = None) -> bytes:
    p = normalize_export_profile(profile)
    buf = BytesIO()
    c = pdf_canvas.Canvas(buf, pagesize=LETTER)
    if resume.template == "modern":
        _draw_modern_pdf(c, resume, p)
    else:
        _draw_classic_pdf(c, resume, p)
    c.save()
    return buf.getvalue()


def build_export_filename(resume: Resume, extension: str) -> str:
    base = clean_text(resume.title).replace(" ", "_") or "resume"
    safe = _FILENAME_PATTERN.sub("", base) or "resume"
    ext  = extension.lstrip(".")
    return f"{safe}.{ext}"


# ===========================================================================
# ── CLASSIC LAYOUT ──────────────────────────────────────────────────────────
#
# Design goals
# ────────────
# • 100 % ATS-parseable: single-column, logical reading order, zero images,
#   zero text boxes, no tables, standard fonts (Helvetica family).
# • Traditional professional look: name centred at top, thin horizontal rule
#   beneath contact bar, bold ALL-CAPS section headings with full-width rule,
#   dates/company right-aligned on same line as role.
# • Tight margins (0.65 ") to maximise content area on one page.
# • Black-on-white throughout – no colour that could confuse optical scanners.
# ===========================================================================

_C_MARGIN    = 0.65 * inch
_C_WIDTH     = PAGE_W - 2 * _C_MARGIN   # 480 pt


def _classic_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    BLK = colors.black
    GRY = colors.HexColor("#444444")
    return {
        # Candidate name
        "name": ParagraphStyle(
            "CName", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=22, leading=26,
            textColor=BLK, alignment=1,   # centred
        ),
        # Job title / tag-line beneath name
        "tagline": ParagraphStyle(
            "CTagline", parent=base["Normal"],
            fontName="Helvetica", fontSize=10.5, leading=13,
            textColor=GRY, alignment=1,
        ),
        # Contact bar (email · phone · location · links)
        "contact": ParagraphStyle(
            "CContact", parent=base["Normal"],
            fontName="Helvetica", fontSize=9.5, leading=12,
            textColor=GRY, alignment=1,
        ),
        # ALL-CAPS section label (Experience, Skills …)
        "section": ParagraphStyle(
            "CSection", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=9.5, leading=11,
            textColor=BLK, spaceAfter=0,
        ),
        # Role / project name  (left-side of two-col header row)
        "entry_role": ParagraphStyle(
            "CEntryRole", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=10.5, leading=13,
            textColor=BLK,
        ),
        # Company name
        "entry_company": ParagraphStyle(
            "CEntryCompany", parent=base["Normal"],
            fontName="Helvetica-Oblique", fontSize=9.8, leading=12,
            textColor=GRY,
        ),
        # Date / location (right-aligned on same row as role)
        "entry_date": ParagraphStyle(
            "CEntryDate", parent=base["Normal"],
            fontName="Helvetica", fontSize=9.5, leading=13,
            textColor=GRY, alignment=2,   # right
        ),
        # Body paragraph (summary)
        "body": ParagraphStyle(
            "CBody", parent=base["Normal"],
            fontName="Helvetica", fontSize=10, leading=14,
            textColor=BLK,
        ),
        # Bullet point
        "bullet": ParagraphStyle(
            "CBullet", parent=base["Normal"],
            fontName="Helvetica", fontSize=10, leading=13.5,
            textColor=BLK,
            leftIndent=10, firstLineIndent=-7,
        ),
        # Skills body
        "skills": ParagraphStyle(
            "CSkills", parent=base["Normal"],
            fontName="Helvetica", fontSize=10, leading=14,
            textColor=BLK,
        ),
    }


def _draw_classic_pdf(
    pdf: pdf_canvas.Canvas,
    resume: Resume,
    profile: dict[str, str],
) -> None:
    styles = _classic_styles()
    y = PAGE_H - 0.55 * inch   # start near top

    # ── Name ────────────────────────────────────────────────────────────────
    display_name = profile["fullName"] or resume.title
    y = _cp(pdf, display_name, styles["name"], _C_MARGIN, y, _C_WIDTH)

    # ── Tag-line ─────────────────────────────────────────────────────────────
    if profile["jobTitle"]:
        y = _cp(pdf, profile["jobTitle"], styles["tagline"],
                _C_MARGIN, y - 2, _C_WIDTH)

    # ── Horizontal rule ───────────────────────────────────────────────────────
    y -= 5
    _hrule(pdf, _C_MARGIN, y, _C_WIDTH, width_pt=1.0)
    y -= 7

    # ── Contact bar ───────────────────────────────────────────────────────────
    contact_parts = [v for v in [
        profile["email"], profile["phone"], profile["location"],
        _strip_scheme(profile["linkedin"]),
        _strip_scheme(profile["github"]),
        _strip_scheme(profile["website"]),
    ] if v]
    if contact_parts:
        y = _cp(pdf, "  ·  ".join(contact_parts), styles["contact"],
                _C_MARGIN, y, _C_WIDTH)

    y -= 10

    # ── Summary ───────────────────────────────────────────────────────────────
    if clean_text(resume.data.summary):
        y = _classic_section_header(pdf, "SUMMARY", y, styles)
        y = _cp(pdf, clean_text(resume.data.summary),
                styles["body"], _C_MARGIN, y - 3, _C_WIDTH)
        y -= 8

    # ── Experience ────────────────────────────────────────────────────────────
    if resume.data.experience:
        y = _classic_section_header(pdf, "EXPERIENCE", y, styles)
        for exp in resume.data.experience:
            # Role (left) + date (right) on one line
            role_txt   = clean_text(exp.role)   or ""
            date_txt   = clean_text(getattr(exp, "date", "")) or ""
            y = _classic_two_col(pdf, role_txt, date_txt,
                                 styles["entry_role"], styles["entry_date"],
                                 _C_MARGIN, y - 3, _C_WIDTH)
            if clean_text(exp.company):
                y = _cp(pdf, clean_text(exp.company),
                        styles["entry_company"], _C_MARGIN, y - 1, _C_WIDTH)
            for pt in exp.points:
                if clean_text(pt):
                    y = _cp(pdf, f"\u2022  {clean_text(pt)}",
                            styles["bullet"], _C_MARGIN, y - 2, _C_WIDTH)
            y -= 5
        y -= 3

    # ── Projects ──────────────────────────────────────────────────────────────
    if resume.data.projects:
        y = _classic_section_header(pdf, "PROJECTS", y, styles)
        for proj in resume.data.projects:
            y = _cp(pdf, clean_text(proj.name),
                    styles["entry_role"], _C_MARGIN, y - 3, _C_WIDTH)
            for pt in proj.points:
                if clean_text(pt):
                    y = _cp(pdf, f"\u2022  {clean_text(pt)}",
                            styles["bullet"], _C_MARGIN, y - 2, _C_WIDTH)
            y -= 5
        y -= 3

    # ── Skills ────────────────────────────────────────────────────────────────
    if resume.data.skills:
        y = _classic_section_header(pdf, "SKILLS", y, styles)
        _cp(pdf, ", ".join(resume.data.skills),
            styles["skills"], _C_MARGIN, y - 3, _C_WIDTH)


def _classic_section_header(
    pdf: pdf_canvas.Canvas,
    title: str,
    y: float,
    styles: dict[str, ParagraphStyle],
) -> float:
    """ALL-CAPS label + full-width 0.75 pt rule underneath."""
    y = _cp(pdf, title, styles["section"], _C_MARGIN, y, _C_WIDTH)
    _hrule(pdf, _C_MARGIN, y + 3, _C_WIDTH, width_pt=0.75)
    return y - 4


def _classic_two_col(
    pdf: pdf_canvas.Canvas,
    left: str,
    right: str,
    left_style: ParagraphStyle,
    right_style: ParagraphStyle,
    x: float,
    y: float,
    total_width: float,
) -> float:
    """Draw left + right text on the same baseline."""
    date_w  = 90   # pt reserved for date column
    role_w  = total_width - date_w - 6

    p_left  = Paragraph(_esc(left),  left_style)
    p_right = Paragraph(_esc(right), right_style)
    _, lh = p_left.wrap(role_w, PAGE_H)
    _, rh = p_right.wrap(date_w, PAGE_H)
    row_h = max(lh, rh)

    if y - row_h < 0.65 * inch:
        pdf.showPage()
        y = PAGE_H - 0.75 * inch

    p_left.drawOn(pdf,  x,                       y - lh)
    p_right.drawOn(pdf, x + role_w + 6,           y - rh)
    return y - row_h


# ===========================================================================
# ── MODERN LAYOUT ────────────────────────────────────────────────────────────
#
# Design goals (2026)
# ───────────────────
# • Sidebar-free single column — still 100 % ATS safe.
# • A full-width colour header band containing name + contact info.
# • Thin accent left-border on section headings (brand colour only).
# • Clean sans-serif with generous leading – readable at a glance.
# • Subtle dot-grid watermark in header (drawn with circles, not images)
#   — picked up by human eyes, invisible to ATS parsers.
# • Colour palette: deep navy #0A192F + electric teal #00C4A0 (2026 tech feel)
# ===========================================================================

_M_MARGIN    = 0.62 * inch
_M_WIDTH     = PAGE_W - 2 * _M_MARGIN
_M_NAVY      = colors.HexColor("#0A192F")
_M_TEAL      = colors.HexColor("#00C4A0")
_M_LIGHT     = colors.HexColor("#F0F4F8")
_M_MUTED     = colors.HexColor("#64748B")
_M_TEXT      = colors.HexColor("#1E293B")
_M_RULE      = colors.HexColor("#CBD5E1")


def _modern_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "MName", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=24, leading=28,
            textColor=colors.white,
        ),
        "tagline": ParagraphStyle(
            "MTagline", parent=base["Normal"],
            fontName="Helvetica", fontSize=11, leading=14,
            textColor=colors.HexColor("#A8BECC"),
        ),
        "contact": ParagraphStyle(
            "MContact", parent=base["Normal"],
            fontName="Helvetica", fontSize=9, leading=12,
            textColor=colors.HexColor("#A8BECC"),
        ),
        "section": ParagraphStyle(
            "MSection", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=9, leading=11,
            textColor=_M_TEAL,
            leftIndent=8,
        ),
        "entry_role": ParagraphStyle(
            "MEntryRole", parent=base["Normal"],
            fontName="Helvetica-Bold", fontSize=11, leading=14,
            textColor=_M_TEXT,
        ),
        "entry_company": ParagraphStyle(
            "MEntryCompany", parent=base["Normal"],
            fontName="Helvetica-Oblique", fontSize=9.8, leading=12,
            textColor=_M_MUTED,
        ),
        "entry_date": ParagraphStyle(
            "MEntryDate", parent=base["Normal"],
            fontName="Helvetica", fontSize=9, leading=14,
            textColor=_M_MUTED, alignment=2,
        ),
        "body": ParagraphStyle(
            "MBody", parent=base["Normal"],
            fontName="Helvetica", fontSize=10, leading=14,
            textColor=_M_TEXT,
        ),
        "bullet": ParagraphStyle(
            "MBullet", parent=base["Normal"],
            fontName="Helvetica", fontSize=10, leading=13.5,
            textColor=_M_TEXT,
            leftIndent=12, firstLineIndent=-8,
        ),
        "skills": ParagraphStyle(
            "MSkills", parent=base["Normal"],
            fontName="Helvetica", fontSize=10, leading=14,
            textColor=_M_TEXT,
        ),
    }


def _draw_modern_pdf(
    pdf: pdf_canvas.Canvas,
    resume: Resume,
    profile: dict[str, str],
) -> None:
    styles = _modern_styles()

    # ── Header band ──────────────────────────────────────────────────────────
    header_h = 1.30 * inch
    pdf.setFillColor(_M_NAVY)
    pdf.rect(0, PAGE_H - header_h, PAGE_W, header_h, stroke=0, fill=1)

    # Subtle dot-grid decoration (top-right corner of band)
    _dot_grid(pdf, PAGE_W - 1.5 * inch, PAGE_H - header_h,
              1.4 * inch, header_h)

    # Teal accent bar at bottom of header
    pdf.setFillColor(_M_TEAL)
    pdf.rect(0, PAGE_H - header_h - 3, PAGE_W, 3, stroke=0, fill=1)

    # Name inside band
    name_x = _M_MARGIN
    name_y = PAGE_H - 0.32 * inch
    p_name = Paragraph(_esc(profile["fullName"] or resume.title), styles["name"])
    _, nh = p_name.wrap(_M_WIDTH * 0.70, header_h)
    name_y -= nh
    p_name.drawOn(pdf, name_x, name_y)

    # Tagline
    if profile["jobTitle"]:
        p_tag = Paragraph(_esc(profile["jobTitle"]), styles["tagline"])
        _, th = p_tag.wrap(_M_WIDTH * 0.70, header_h)
        name_y -= th + 2
        p_tag.drawOn(pdf, name_x, name_y)

    # Contact info (right-aligned inside band)
    contact_parts = [v for v in [
        profile["email"], profile["phone"], profile["location"],
        _strip_scheme(profile["linkedin"]),
        _strip_scheme(profile["github"]),
        _strip_scheme(profile["website"]),
    ] if v]
    if contact_parts:
        contact_block_x = PAGE_W * 0.55
        contact_block_w = PAGE_W - contact_block_x - _M_MARGIN
        cy = PAGE_H - 0.34 * inch
        for part in contact_parts:
            p = Paragraph(_esc(part), styles["contact"])
            _, ph = p.wrap(contact_block_w, header_h)
            cy -= ph
            p.drawOn(pdf, contact_block_x, cy)
            cy -= 2

    # ── Body area ────────────────────────────────────────────────────────────
    y = PAGE_H - header_h - 3 - 14   # 14 pt gap below teal stripe

    # ── Summary ───────────────────────────────────────────────────────────────
    if clean_text(resume.data.summary):
        y = _modern_section_header(pdf, "SUMMARY", y, styles)
        y = _cp(pdf, clean_text(resume.data.summary),
                styles["body"], _M_MARGIN, y - 4, _M_WIDTH)
        y -= 10

    # ── Experience ────────────────────────────────────────────────────────────
    if resume.data.experience:
        y = _modern_section_header(pdf, "EXPERIENCE", y, styles)
        for exp in resume.data.experience:
            date_txt = clean_text(getattr(exp, "date", "")) or ""
            role_txt = clean_text(exp.role) or ""
            y = _modern_entry_header(pdf, role_txt, date_txt,
                                     styles, _M_MARGIN, y - 4, _M_WIDTH)
            if clean_text(exp.company):
                y = _cp(pdf, clean_text(exp.company),
                        styles["entry_company"], _M_MARGIN, y - 1, _M_WIDTH)
            for pt in exp.points:
                if clean_text(pt):
                    y = _cp(pdf, f"\u2022  {clean_text(pt)}",
                            styles["bullet"], _M_MARGIN, y - 2, _M_WIDTH)
            y -= 6
        y -= 2

    # ── Projects ──────────────────────────────────────────────────────────────
    if resume.data.projects:
        y = _modern_section_header(pdf, "PROJECTS", y, styles)
        for proj in resume.data.projects:
            y = _cp(pdf, clean_text(proj.name),
                    styles["entry_role"], _M_MARGIN, y - 4, _M_WIDTH)
            for pt in proj.points:
                if clean_text(pt):
                    y = _cp(pdf, f"\u2022  {clean_text(pt)}",
                            styles["bullet"], _M_MARGIN, y - 2, _M_WIDTH)
            y -= 6
        y -= 2

    # ── Skills ────────────────────────────────────────────────────────────────
    if resume.data.skills:
        y = _modern_section_header(pdf, "SKILLS", y, styles)
        _cp(pdf, ", ".join(resume.data.skills),
            styles["skills"], _M_MARGIN, y - 4, _M_WIDTH)


def _modern_section_header(
    pdf: pdf_canvas.Canvas,
    title: str,
    y: float,
    styles: dict[str, ParagraphStyle],
) -> float:
    """Teal left accent bar + ALL-CAPS label + light horizontal rule."""
    # Teal left accent
    bar_h = 13
    pdf.setFillColor(_M_TEAL)
    pdf.rect(_M_MARGIN, y - bar_h + 4, 3, bar_h, stroke=0, fill=1)

    y = _cp(pdf, title, styles["section"],
            _M_MARGIN, y, _M_WIDTH)

    # Full-width light rule
    pdf.setStrokeColor(_M_RULE)
    pdf.setLineWidth(0.5)
    pdf.line(_M_MARGIN, y + 3, _M_MARGIN + _M_WIDTH, y + 3)
    return y - 5


def _modern_entry_header(
    pdf: pdf_canvas.Canvas,
    role: str,
    date: str,
    styles: dict[str, ParagraphStyle],
    x: float,
    y: float,
    total_width: float,
) -> float:
    date_w = 100
    role_w = total_width - date_w - 8
    p_role = Paragraph(_esc(role), styles["entry_role"])
    p_date = Paragraph(_esc(date), styles["entry_date"])
    _, rh = p_role.wrap(role_w, PAGE_H)
    _, dh = p_date.wrap(date_w, PAGE_H)
    row_h = max(rh, dh)
    if y - row_h < 0.62 * inch:
        pdf.showPage()
        y = PAGE_H - 0.72 * inch
    p_role.drawOn(pdf, x,             y - rh)
    p_date.drawOn(pdf, x + role_w + 8, y - dh)
    return y - row_h


# ===========================================================================
# Low-level drawing utilities
# ===========================================================================

def _cp(
    pdf: pdf_canvas.Canvas,
    text: str,
    style: ParagraphStyle,
    x: float,
    y: float,
    width: float,
    centered: bool = False,
) -> float:
    """Draw a Paragraph and return new y (below paragraph)."""
    if not text:
        return y
    p = Paragraph(_esc(text), style)
    _, h = p.wrap(width, PAGE_H)
    if y - h < 0.62 * inch:
        pdf.showPage()
        y = PAGE_H - 0.75 * inch
    draw_x = ((PAGE_W - width) / 2) if centered else x
    p.drawOn(pdf, draw_x, y - h)
    return y - h


def _hrule(
    pdf: pdf_canvas.Canvas,
    x: float,
    y: float,
    width: float,
    width_pt: float = 0.75,
    colour: colors.Color = colors.black,
) -> None:
    pdf.setStrokeColor(colour)
    pdf.setLineWidth(width_pt)
    pdf.line(x, y, x + width, y)


def _dot_grid(
    pdf: pdf_canvas.Canvas,
    x: float,
    y: float,
    w: float,
    h: float,
    step: float = 14,
    radius: float = 1.2,
) -> None:
    """Render a subtle dot-grid in the header corner (decorative only)."""
    pdf.setFillColor(colors.HexColor("#1B3A5C"))
    cols = int(w / step) + 1
    rows = int(h / step) + 1
    for r in range(rows):
        for c in range(cols):
            pdf.circle(x + c * step, y + r * step, radius, stroke=0, fill=1)


def _strip_scheme(url: str) -> str:
    if not url:
        return ""
    return url.replace("https://", "").replace("http://", "").rstrip("/")


def _esc(text: str) -> str:
    """Escape XML special chars for ReportLab Paragraph."""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )